#!/usr/bin/env python3
"""
Generate three user manuals (Word .docx) for MAXEVA GESTIÓN ERP:
  1. Manual del Administrador
  2. Manual del Contable
  3. Manual del Vendedor
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "docs", "manuales")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ─── Styling helpers ─────────────────────────────────────────────────────────

def set_doc_defaults(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(2)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hf = hs.font
        hf.name = "Calibri"
        hf.bold = True
        if level == 1:
            hf.size = Pt(22)
            hf.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        elif level == 2:
            hf.size = Pt(16)
            hf.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        else:
            hf.size = Pt(13)
            hf.color.rgb = RGBColor(0x37, 0x4, 0x9B)


def add_cover(doc, title, subtitle, role_desc):
    for _ in range(6):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph("")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(role_desc)
    run3.font.size = Pt(13)
    run3.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    run3.italic = True

    for _ in range(4):
        doc.add_paragraph("")

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Maxeva Gestión — Sistema de Gestión para Agencias de Viaje")
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("Versión 1.0 — Marzo 2026")
    run5.font.size = Pt(10)
    run5.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    doc.add_page_break()


def h1(doc, text):
    doc.add_heading(text, level=1)

def h2(doc, text):
    doc.add_heading(text, level=2)

def h3(doc, text):
    doc.add_heading(text, level=3)

def p(doc, text):
    doc.add_paragraph(text)

def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")

def numbered(doc, text):
    doc.add_paragraph(text, style="List Number")

def bold_p(doc, bold_text, normal_text=""):
    para = doc.add_paragraph()
    run = para.add_run(bold_text)
    run.bold = True
    if normal_text:
        para.add_run(normal_text)

def tip(doc, text):
    para = doc.add_paragraph()
    run = para.add_run("Consejo: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
    para.add_run(text)

def important(doc, text):
    para = doc.add_paragraph()
    run = para.add_run("Importante: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    para.add_run(text)

def pagebreak(doc):
    doc.add_page_break()


# ─── Shared Sections ─────────────────────────────────────────────────────────

def add_intro(doc):
    h1(doc, "Introducción")
    p(doc, "Maxeva Gestión es el sistema integral para la gestión de tu agencia de viajes. Desde este sistema podés administrar todo el flujo comercial: desde que llega una consulta de un cliente hasta que se cierra la operación, se cobran los pagos y se calculan las comisiones.")
    p(doc, "Este manual te va a guiar paso a paso por cada sección del sistema, explicando qué podés hacer, cómo hacerlo, y cómo las acciones en un lugar impactan en otros.")

def add_login(doc):
    h1(doc, "Acceso al Sistema")
    p(doc, "Para ingresar al sistema:")
    numbered(doc, "Abrí tu navegador y entrá a la dirección del sistema (ej: maxevagestion.com)")
    numbered(doc, "Ingresá tu email y contraseña")
    numbered(doc, "Hacé clic en \"Iniciar Sesión\"")
    p(doc, "Si olvidaste tu contraseña, hacé clic en \"Olvidé mi contraseña\" y seguí las instrucciones que te llegarán por email.")
    important(doc, "No compartas tus credenciales con nadie. Cada usuario tiene permisos diferentes según su rol.")


# ─── ADMIN MANUAL ─────────────────────────────────────────────────────────────

def generate_admin_manual():
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc, "MAXEVA GESTIÓN", "Manual del Administrador", "Guía completa para la gestión integral del sistema")

    # TOC placeholder
    h1(doc, "Contenido")
    p(doc, "1. Introducción")
    p(doc, "2. Acceso al Sistema")
    p(doc, "3. Panel Principal (Dashboard)")
    p(doc, "4. CRM de Ventas — Gestión de Leads")
    p(doc, "5. Cotizaciones")
    p(doc, "6. Operaciones")
    p(doc, "7. Clientes")
    p(doc, "8. Operadores / Proveedores")
    p(doc, "9. Caja y Bancos")
    p(doc, "10. Contabilidad")
    p(doc, "11. Comisiones")
    p(doc, "12. Reportes")
    p(doc, "13. Alertas")
    p(doc, "14. Herramientas")
    p(doc, "15. Configuración")
    p(doc, "16. Flujo Completo: De Lead a Cierre")
    pagebreak(doc)

    add_intro(doc)
    p(doc, "Como Administrador, tenés acceso completo a todas las secciones del sistema. Podés ver toda la información de la agencia, crear y editar registros, y gestionar el equipo de trabajo.")
    pagebreak(doc)

    add_login(doc)
    pagebreak(doc)

    # ─── Dashboard ───
    h1(doc, "Panel Principal (Dashboard)")
    p(doc, "El Dashboard es la primera pantalla que ves al entrar. Te muestra un resumen en tiempo real de cómo va la agencia.")

    h2(doc, "Indicadores Principales (KPIs)")
    p(doc, "En la parte superior vas a encontrar los números clave:")
    bullet(doc, "Total de ventas del período (en pesos y en dólares)")
    bullet(doc, "Cantidad de operaciones activas")
    bullet(doc, "Pagos pendientes de cobro (lo que te deben los clientes)")
    bullet(doc, "Pagos pendientes a operadores (lo que vos le debés a los proveedores)")
    bullet(doc, "Comisiones pendientes de pago a vendedores")

    h2(doc, "Filtros")
    p(doc, "Podés filtrar toda la información por:")
    bullet(doc, "Rango de fechas (últimos 7 días, 30 días, mes actual, personalizado)")
    bullet(doc, "Agencia (si manejás varias sucursales)")
    bullet(doc, "Vendedor (para ver los números de un vendedor en particular)")

    h2(doc, "Gráficos")
    p(doc, "Debajo de los KPIs encontrás gráficos interactivos:")
    bullet(doc, "Ventas por destino: para saber cuáles son los destinos más vendidos")
    bullet(doc, "Ventas por vendedor: para comparar el rendimiento del equipo")
    bullet(doc, "Análisis de margen: cuánto ganás realmente en cada operación")

    tip(doc, "Usá los filtros para analizar períodos específicos. Por ejemplo, compará las ventas del mes pasado con el anterior para ver la tendencia.")

    bold_p(doc, "Impacto: ", "los números del Dashboard se actualizan automáticamente cuando se crean operaciones, se registran pagos o cambian estados. No necesitás hacer nada manual.")
    pagebreak(doc)

    # ─── CRM ───
    h1(doc, "CRM de Ventas — Gestión de Leads")
    p(doc, "El CRM es donde gestionás todas las consultas y oportunidades de venta. Un \"lead\" es un potencial cliente que mostró interés en un viaje.")

    h2(doc, "Cómo llegan los leads")
    p(doc, "Los leads pueden llegar de tres formas:")
    numbered(doc, "Manual: los creás vos directamente desde el botón \"Nuevo Lead\"")
    numbered(doc, "Desde Trello: si tenés la integración activa, cuando alguien crea una tarjeta en Trello, aparece automáticamente acá")
    numbered(doc, "Desde Manychat/WhatsApp: los contactos nuevos que llegan por Manychat se crean como leads automáticamente")

    h2(doc, "Vista Kanban y Vista Tabla")
    p(doc, "Podés ver los leads de dos formas:")
    bullet(doc, "Kanban: tablero con columnas por estado (Nuevo, En Progreso, Cotizado, Ganado, Perdido). Arrastrás las tarjetas para cambiar el estado.")
    bullet(doc, "Tabla: lista ordenable y filtrable con todos los datos")

    h2(doc, "Estados de un Lead")
    bullet(doc, "Nuevo: acaba de llegar la consulta")
    bullet(doc, "En Progreso: estás trabajando en la consulta, buscando opciones")
    bullet(doc, "Cotizado: ya le enviaste una cotización al cliente")
    bullet(doc, "Ganado: el cliente confirmó y se convierte en una operación")
    bullet(doc, "Perdido: el cliente no avanzó")

    h2(doc, "Crear un Lead Nuevo")
    numbered(doc, "Hacé clic en \"Nuevo Lead\"")
    numbered(doc, "Completá el nombre del contacto, teléfono y/o email")
    numbered(doc, "Indicá el destino de interés y la región")
    numbered(doc, "Asignalo a un vendedor (o dejalo sin asignar)")
    numbered(doc, "Guardá")

    h2(doc, "Detalle de un Lead")
    p(doc, "Al hacer clic en un lead, se abre su ficha completa donde podés:")
    bullet(doc, "Ver y editar toda la información del contacto")
    bullet(doc, "Ver el historial de cotizaciones")
    bullet(doc, "Crear una nueva cotización")
    bullet(doc, "Convertirlo en operación (cuando el cliente confirma)")
    bullet(doc, "Enviar un WhatsApp directamente")

    bold_p(doc, "Impacto: ", "cuando cambiás el estado de un lead a \"Ganado\" y lo convertís en operación, automáticamente se crea un registro de operación, se crea el cliente (si no existía), y si ya tenía una seña/depósito registrado, ese movimiento se transfiere a la operación.")

    important(doc, "Si tenés Trello integrado, al mover un lead de estado o asignarlo a un vendedor, la tarjeta en Trello también se mueve automáticamente.")
    pagebreak(doc)

    # ─── Quotations ───
    h1(doc, "Cotizaciones")
    p(doc, "Las cotizaciones te permiten armar propuestas detalladas para tus clientes y compartirlas con un link público.")

    h2(doc, "Crear una Cotización")
    numbered(doc, "Entrá al detalle de un lead")
    numbered(doc, "Hacé clic en \"Cotizar\"")
    numbered(doc, "Se abre el armador de cotización donde podés agregar distintos servicios:")

    h3(doc, "Tipos de Servicio")
    bullet(doc, "Hotel: buscá hoteles por nombre — el sistema tiene una base de más de 1.600 hoteles con fotos, estrellas y dirección. Al seleccionar uno, se completan automáticamente los datos.")
    bullet(doc, "Vuelo: indicá aerolínea, ruta, fechas de ida y vuelta, clase y escalas")
    bullet(doc, "Asistencia al viajero: proveedor, tipo de cobertura y precio")
    bullet(doc, "Traslado: descripción del transfer y precio")
    bullet(doc, "Paquete: para combinar varios servicios en uno solo")
    bullet(doc, "Otro: cualquier servicio adicional")

    h2(doc, "Opciones Múltiples")
    p(doc, "Podés armar hasta 3 opciones diferentes dentro de la misma cotización. Por ejemplo: Opción 1 con hotel 4 estrellas y Opción 2 con hotel 5 estrellas, para que el cliente compare.")

    h2(doc, "Link Público")
    p(doc, "Una vez que guardás la cotización, el sistema genera un link que podés compartir por WhatsApp o email. El cliente abre ese link y ve una página profesional con todos los detalles del viaje, fotos de hoteles y precios.")

    h2(doc, "Estados de la Cotización")
    bullet(doc, "Borrador: todavía estás armándola, podés editarla")
    bullet(doc, "Enviada: ya la compartiste con el cliente")
    bullet(doc, "Convertida: el cliente aceptó y se creó la operación")

    bold_p(doc, "Impacto: ", "las cotizaciones quedan asociadas al lead. Si el lead se convierte en operación, la cotización queda vinculada para referencia futura.")
    pagebreak(doc)

    # ─── Operations ───
    h1(doc, "Operaciones")
    p(doc, "Una operación es un viaje confirmado. Es el corazón del sistema — acá se registra toda la información del viaje vendido.")

    h2(doc, "Crear una Operación")
    p(doc, "Las operaciones se pueden crear de dos formas:")
    numbered(doc, "Desde un Lead: al convertir un lead ganado, se crea automáticamente con los datos del contacto")
    numbered(doc, "Manual: desde el botón \"Nueva Operación\" en la sección de Operaciones")

    h2(doc, "Datos de la Operación")
    bullet(doc, "Tipo: Vuelo, Hotel, Paquete, Crucero, Traslado, Asistencia o Mixto")
    bullet(doc, "Destino y fechas de salida/regreso")
    bullet(doc, "Agencia asignada")
    bullet(doc, "Vendedor principal (y opcionalmente un vendedor secundario)")
    bullet(doc, "Monto de venta total")
    bullet(doc, "Moneda (pesos o dólares)")

    h2(doc, "Clientes de la Operación")
    p(doc, "Una operación puede tener varios pasajeros/clientes asociados. El sistema lleva el dato de quién es el titular y quiénes son los acompañantes.")

    h2(doc, "Operadores/Proveedores")
    p(doc, "Cada operación tiene uno o más operadores (el hotel, la aerolínea, el proveedor de asistencia). Para cada uno se registra:")
    bullet(doc, "El costo que cobrá el operador")
    bullet(doc, "La moneda del costo")
    bullet(doc, "El estado de pago al operador")

    h2(doc, "Estados de la Operación")
    bullet(doc, "Pre-reserva: el viaje está en proceso de reserva")
    bullet(doc, "Reservado: confirmada la reserva con el operador")
    bullet(doc, "Confirmado: todo pago, viaje confirmado")
    bullet(doc, "Viajado: el cliente ya viajó")
    bullet(doc, "Cerrado: operación finalizada, todo liquidado")

    h2(doc, "Margen de Ganancia")
    p(doc, "El sistema calcula automáticamente el margen restando el costo del operador al precio de venta. Este margen es la base para calcular las comisiones de los vendedores.")

    bold_p(doc, "Impactos de las operaciones:")
    bullet(doc, "Cuando una operación pasa a \"Confirmado\", se calculan automáticamente las comisiones del vendedor")
    bullet(doc, "Los pagos que se registran en la operación generan movimientos contables automáticos")
    bullet(doc, "Los saldos de los operadores se actualizan con cada pago registrado")
    bullet(doc, "Las alertas se generan automáticamente (viaje próximo, pago pendiente, documentación faltante)")

    important(doc, "Si cambiás el monto de venta o el costo del operador después de haber calculado comisiones, las comisiones NO se recalculan automáticamente. Tendrás que recalcularlas desde la sección de Comisiones.")
    pagebreak(doc)

    # ─── Customers ───
    h1(doc, "Clientes")
    p(doc, "La sección de Clientes guarda toda la información de las personas que viajan con tu agencia.")

    h2(doc, "Formas de Crear un Cliente")
    numbered(doc, "Automático desde un Lead: cuando convertís un lead en operación, el sistema crea el cliente con los datos del contacto")
    numbered(doc, "Automático desde OCR: cuando subís un pasaporte o DNI, el sistema lee los datos y crea o actualiza el cliente")
    numbered(doc, "Manual: desde el botón \"Nuevo Cliente\"")

    h2(doc, "Datos del Cliente")
    bullet(doc, "Nombre y apellido")
    bullet(doc, "Email y teléfono")
    bullet(doc, "Fecha de nacimiento")
    bullet(doc, "Tipo y número de documento (DNI, Pasaporte)")
    bullet(doc, "Nacionalidad")
    bullet(doc, "Notas adicionales")

    h2(doc, "Historial del Cliente")
    p(doc, "Desde la ficha de cada cliente podés ver:")
    bullet(doc, "Todas las operaciones en las que participó")
    bullet(doc, "El total que gastó (en pesos y dólares)")
    bullet(doc, "Los pagos realizados")

    h2(doc, "Lectura Automática de Documentos (OCR)")
    p(doc, "Podés subir una foto del pasaporte o DNI del cliente, y el sistema lee automáticamente:")
    bullet(doc, "Nombre completo")
    bullet(doc, "Número de documento")
    bullet(doc, "Fecha de nacimiento")
    bullet(doc, "Fecha de vencimiento del documento")
    p(doc, "El sistema te muestra los datos extraídos para que los confirmes antes de guardar.")

    bold_p(doc, "Impacto: ", "los clientes se vinculan a las operaciones. Un mismo cliente puede tener múltiples viajes, y desde su ficha podés ver todo el historial.")
    pagebreak(doc)

    # ─── Operators ───
    h1(doc, "Operadores / Proveedores")
    p(doc, "Los operadores son los proveedores con los que trabajás: hoteles, aerolíneas, agencias mayoristas, empresas de asistencia, etc.")

    h2(doc, "Datos del Operador")
    bullet(doc, "Nombre de la empresa")
    bullet(doc, "Contacto (nombre, email, teléfono)")
    bullet(doc, "Datos bancarios (para pagos)")
    bullet(doc, "Límite de crédito")
    bullet(doc, "Notas")

    h2(doc, "Balance del Operador")
    p(doc, "El sistema lleva automáticamente el balance de cada operador:")
    bullet(doc, "Cuánto le debés en total (sumando todas las operaciones)")
    bullet(doc, "Cuánto ya le pagaste")
    bullet(doc, "Cuánto falta pagar")
    bullet(doc, "Próximo vencimiento de pago")

    bold_p(doc, "Impacto: ", "cada vez que registrás un pago a un operador en una operación, el balance general de ese operador se actualiza automáticamente. También se genera un asiento contable.")
    pagebreak(doc)

    # ─── Cash ───
    h1(doc, "Caja y Bancos")
    p(doc, "Esta sección te muestra el estado financiero en tiempo real: cuánto dinero hay disponible, qué pagos entraron y qué salió.")

    h2(doc, "Pestañas")
    bullet(doc, "Resumen: saldo actual por moneda (pesos, dólares) y por método (efectivo, banco, Mercado Pago)")
    bullet(doc, "Caja: movimientos de efectivo")
    bullet(doc, "Pagos Recibidos: cobros de clientes")
    bullet(doc, "Movimientos: todos los movimientos (ingresos y egresos)")
    bullet(doc, "Egresos: pagos realizados (a operadores, comisiones, gastos)")
    bullet(doc, "Ingresos: cobros de clientes y otros ingresos")

    h2(doc, "Registrar un Pago de Cliente")
    numbered(doc, "Podés hacerlo desde la operación (en la pestaña de pagos) o desde Caja")
    numbered(doc, "Indicá monto, moneda, método de pago y fecha")
    numbered(doc, "Guardá el pago")

    h2(doc, "Registrar un Pago a Operador")
    numbered(doc, "Desde la operación o desde la ficha del operador")
    numbered(doc, "Indicá monto, moneda, método y número de comprobante")
    numbered(doc, "Al marcarlo como \"Pagado\" se genera el movimiento contable")

    bold_p(doc, "Impactos cuando se marca un pago como \"Pagado\":")
    bullet(doc, "Se crea automáticamente un asiento en el libro contable (doble entrada)")
    bullet(doc, "Se actualiza el saldo de caja")
    bullet(doc, "Se actualiza el balance del operador (si es un pago a proveedor)")
    bullet(doc, "Si es un cobro de cliente, se actualiza el estado de la operación")

    important(doc, "Los movimientos contables se crean automáticamente. No hace falta cargarlos a mano en Contabilidad — todo parte de los pagos.")
    pagebreak(doc)

    # ─── Accounting ───
    h1(doc, "Contabilidad")
    p(doc, "La sección contable muestra todos los movimientos financieros organizados por partida doble. La mayoría de los asientos se generan automáticamente a partir de los pagos.")

    h2(doc, "Libro Mayor (Ledger)")
    p(doc, "Es el registro central de todos los movimientos contables. Cada transacción genera dos movimientos: uno de débito y uno de crédito (partida doble).")
    p(doc, "Podés filtrar por: cuenta, fecha, operación, vendedor o tipo de movimiento.")

    h2(doc, "Posición IVA")
    p(doc, "Muestra el IVA fiscal (ventas) vs. el IVA crédito (compras) y la posición mensual. Esto te ayuda a saber cuánto IVA tenés que pagar.")

    h2(doc, "Pagos a Operadores")
    p(doc, "Vista consolidada de todos los pagos pendientes y realizados a proveedores. Desde acá podés ver quiénes necesitan pago urgente.")

    h2(doc, "Cuentas Corrientes")
    p(doc, "Balance detallado por operador y por cliente, mostrando antigüedad de deuda (30, 60, 90+ días).")

    h2(doc, "Multimoneda")
    p(doc, "El sistema maneja pesos y dólares. Todas las transacciones en dólares también se registran con su equivalente en pesos al tipo de cambio del momento. Si hay diferencias de cambio, se registran automáticamente como ganancia o pérdida por tipo de cambio.")

    bold_p(doc, "Impacto: ", "la contabilidad se alimenta automáticamente de los pagos. Cuando un pago se marca como cobrado/pagado, el asiento se crea solo. Esto significa que la contabilidad siempre refleja el estado real del negocio sin carga manual.")
    pagebreak(doc)

    # ─── Commissions ───
    h1(doc, "Comisiones")
    p(doc, "El sistema calcula automáticamente las comisiones de cada vendedor cuando una operación llega al estado \"Confirmado\" o \"Cerrado\".")

    h2(doc, "Cómo se Calculan")
    p(doc, "La fórmula es simple:")
    bold_p(doc, "Margen", " = Precio de Venta - Costo del Operador")
    bold_p(doc, "Comisión", " = Margen × Porcentaje del Vendedor")
    p(doc, "El porcentaje de cada vendedor se configura en las Reglas de Comisión.")

    h2(doc, "Vendedor Secundario")
    p(doc, "Si una operación tiene un vendedor principal y uno secundario, la comisión se divide entre los dos (por defecto 50/50, pero es configurable).")

    h2(doc, "Vista de Comisiones")
    p(doc, "Desde esta sección podés:")
    bullet(doc, "Ver todas las comisiones pendientes de pago")
    bullet(doc, "Ver el detalle de cálculo de cada una")
    bullet(doc, "Marcar comisiones como pagadas (individual o en lote)")
    bullet(doc, "Ver el historial de pagos de comisiones por vendedor")
    bullet(doc, "Exportar el reporte de comisiones")

    bold_p(doc, "Impacto: ", "cuando marcás una comisión como pagada, se genera un asiento contable y se descuenta del saldo de caja.")
    pagebreak(doc)

    # ─── Reports ───
    h1(doc, "Reportes")
    p(doc, "La sección de reportes te permite analizar la información del negocio desde distintos ángulos.")

    h2(doc, "Reportes Disponibles")
    bullet(doc, "Reporte de Ventas: total vendido por vendedor, destino o período")
    bullet(doc, "Reporte de Rentabilidad: margen real por operación y por vendedor")
    bullet(doc, "Reporte de Flujo de Caja: ingresos vs. egresos por período")
    bullet(doc, "Balance de Operadores: lo que debés a cada proveedor")
    bullet(doc, "Antigüedad de Deuda: clientes que deben dinero, agrupados por 30/60/90+ días")
    bullet(doc, "Reporte de Comisiones: cuánto ganó cada vendedor")

    h2(doc, "Exportación")
    p(doc, "Todos los reportes se pueden exportar a CSV/Excel para trabajarlos en una planilla.")

    tip(doc, "Usá los reportes periódicamente para detectar tendencias. Por ejemplo, el reporte de rentabilidad te muestra si estás vendiendo con buen margen o si necesitás renegociar costos con algún operador.")
    pagebreak(doc)

    # ─── Alerts ───
    h1(doc, "Alertas")
    p(doc, "El sistema genera alertas automáticas para que no se te pase nada importante.")

    h2(doc, "Tipos de Alertas")
    bullet(doc, "Pago pendiente de cliente: cuando se acerca un vencimiento de cobro")
    bullet(doc, "Pago pendiente a operador: cuando hay que pagarle a un proveedor")
    bullet(doc, "Viaje próximo: 48 a 72 horas antes de la fecha de salida de un viaje")
    bullet(doc, "Documentación faltante: cuando un pasajero no tiene documentos cargados")
    bullet(doc, "Saldo bajo en caja: si el efectivo disponible baja de un umbral")
    bullet(doc, "Vencimiento de IVA: recordatorio de fechas impositivas")
    bullet(doc, "Pérdida por tipo de cambio: cuando hay diferencias significativas de cambio")

    p(doc, "Las alertas se pueden marcar como leídas, archivar o posponer.")
    pagebreak(doc)

    # ─── Tools ───
    h1(doc, "Herramientas")

    h2(doc, "Cerebro (Asistente IA)")
    p(doc, "Es un asistente inteligente al que le podés hacer preguntas sobre tu negocio en lenguaje natural. Ejemplos:")
    bullet(doc, "\"¿Cuánto vendí este mes en dólares?\"")
    bullet(doc, "\"¿Cuál es el operador que más me debe?\"")
    bullet(doc, "\"¿Quién es mi mejor vendedor?\"")
    bullet(doc, "\"¿Cuáles son los destinos más populares?\"")

    h2(doc, "Emilia (Buscador de Viajes)")
    p(doc, "Herramienta para buscar vuelos, hoteles y paquetes en tiempo real. Los resultados se pueden usar para armar cotizaciones directamente.")

    h2(doc, "Tareas")
    p(doc, "Sistema de tareas para el equipo. Podés crear tareas, asignarlas a miembros del equipo, poner fecha de vencimiento y marcarlas como completadas.")

    h2(doc, "Mensajes WhatsApp")
    p(doc, "Desde acá podés enviar mensajes a clientes por WhatsApp usando plantillas predefinidas. Las plantillas pueden incluir el nombre del cliente, destino, fecha de viaje, etc.")

    h2(doc, "Calendario")
    p(doc, "Vista de calendario con todas las fechas de salida de viajes, vencimientos de pagos y eventos del equipo.")
    pagebreak(doc)

    # ─── Settings ───
    h1(doc, "Configuración")
    p(doc, "Desde Configuración podés administrar el sistema completo.")

    h2(doc, "Usuarios")
    p(doc, "Crear, editar y desactivar usuarios del sistema. Cada usuario tiene un rol que determina qué puede ver y hacer:")
    bullet(doc, "Administrador: acceso completo (como vos)")
    bullet(doc, "Vendedor: solo ve sus propios leads, operaciones y comisiones")
    bullet(doc, "Contable: solo ve la parte financiera (caja, contabilidad, operadores)")
    bullet(doc, "Visualizador: ve todo pero no puede modificar nada")

    h2(doc, "Equipos")
    p(doc, "Agrupá vendedores en equipos para organizarlos mejor.")

    h2(doc, "Integraciones")
    bullet(doc, "Trello: configurá la conexión con tu tablero de Trello para que los leads se sincronicen automáticamente")
    bullet(doc, "WhatsApp/Manychat: configurá la conexión para recibir contactos de Manychat")
    bullet(doc, "Email: configurá las notificaciones por email")

    h2(doc, "Importación")
    p(doc, "Podés importar clientes, operadores y operaciones en lote usando archivos CSV.")

    h2(doc, "Reglas de Comisión")
    p(doc, "Configurá el porcentaje de comisión para cada vendedor. Podés crear reglas por vendedor específico o una regla genérica que aplique a todos.")
    pagebreak(doc)

    # ─── Complete flow ───
    h1(doc, "Flujo Completo: De Lead a Cierre")
    p(doc, "Este es el recorrido completo de una venta en el sistema, y cómo cada paso impacta en otros:")

    h2(doc, "1. Llega una consulta (Lead)")
    p(doc, "Un potencial cliente pregunta por un viaje. Se crea un lead en el CRM (manual, desde Trello, o desde Manychat).")
    bold_p(doc, "Impacta en: ", "CRM (aparece como lead nuevo), Dashboard (cuenta como lead activo)")

    h2(doc, "2. Se arma una cotización")
    p(doc, "El vendedor arma una cotización con hoteles, vuelos y servicios, y le comparte el link al cliente.")
    bold_p(doc, "Impacta en: ", "el lead pasa a estado \"Cotizado\"")

    h2(doc, "3. El cliente confirma")
    p(doc, "Se convierte el lead en operación.")
    bold_p(doc, "Impacta en: ", "se crea la Operación, se crea el Cliente (si no existía), el lead pasa a \"Ganado\", aparece en el Dashboard como venta")

    h2(doc, "4. Se registran los pagos")
    p(doc, "Se cargan los cobros del cliente y los pagos a operadores.")
    bold_p(doc, "Impacta en: ", "Caja (se actualiza el saldo), Contabilidad (se crean asientos automáticos), Balance del Operador (se actualiza)")

    h2(doc, "5. Se confirma la operación")
    p(doc, "La operación pasa a estado \"Confirmado\".")
    bold_p(doc, "Impacta en: ", "Comisiones (se calcula automáticamente cuánto le corresponde al vendedor), Alertas (se activan recordatorios de viaje)")

    h2(doc, "6. El cliente viaja")
    p(doc, "Se actualiza el estado a \"Viajado\".")

    h2(doc, "7. Se cierra la operación")
    p(doc, "Se liquidan todos los pagos, se pagan las comisiones.")
    bold_p(doc, "Impacta en: ", "Comisiones (se marcan como pagadas), Contabilidad (asientos de comisión), Caja (se descuenta el pago de comisión), Reportes (se refleja en todos los informes)")

    doc.save(os.path.join(OUTPUT_DIR, "Manual_Administrador_MaxevaGestion.docx"))
    print("✅ Manual del Administrador generado")


# ─── CONTABLE MANUAL ─────────────────────────────────────────────────────────

def generate_contable_manual():
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc, "MAXEVA GESTIÓN", "Manual del Contable", "Guía para la gestión financiera y contable del sistema")

    h1(doc, "Contenido")
    p(doc, "1. Introducción")
    p(doc, "2. Acceso al Sistema")
    p(doc, "3. Tu Menú — Qué Podés Ver y Qué No")
    p(doc, "4. Operaciones (Vista de Lectura)")
    p(doc, "5. Operadores / Proveedores")
    p(doc, "6. Caja y Bancos")
    p(doc, "7. Contabilidad")
    p(doc, "8. Comisiones (Vista de Lectura)")
    p(doc, "9. Reportes")
    p(doc, "10. Alertas")
    p(doc, "11. Tareas")
    p(doc, "12. Cómo las Acciones se Conectan Entre Sí")
    pagebreak(doc)

    add_intro(doc)
    p(doc, "Como Contable, tu rol se enfoca en la parte financiera del negocio: controlar los pagos, verificar la contabilidad, monitorear los saldos de operadores y revisar las comisiones.")
    p(doc, "No tenés acceso a las secciones comerciales (Dashboard, Leads, Clientes) ni a la Configuración del sistema. Esto es para que puedas concentrarte en lo financiero sin distracciones.")
    pagebreak(doc)

    add_login(doc)
    pagebreak(doc)

    # ─── What you see ───
    h1(doc, "Tu Menú — Qué Podés Ver y Qué No")

    h2(doc, "Secciones a las que tenés acceso")
    bullet(doc, "Operaciones: podés ver las operaciones (solo lectura, no podés modificarlas)")
    bullet(doc, "Operadores: podés ver y editar datos de proveedores")
    bullet(doc, "Caja y Bancos: acceso completo para registrar y consultar movimientos")
    bullet(doc, "Contabilidad: acceso completo al libro mayor, IVA, retenciones y toda la parte impositiva")
    bullet(doc, "Comisiones: podés ver las comisiones calculadas (solo lectura)")
    bullet(doc, "Reportes: podés generar y consultar reportes financieros")
    bullet(doc, "Alertas: ves las alertas del sistema relacionadas con vencimientos y finanzas")
    bullet(doc, "Tareas: podés gestionar tus tareas asignadas")

    h2(doc, "Secciones que NO ves")
    bullet(doc, "Dashboard general")
    bullet(doc, "CRM / Leads")
    bullet(doc, "Clientes")
    bullet(doc, "Documentos")
    bullet(doc, "Configuración")
    p(doc, "Estas secciones son de uso comercial o administrativo y no aparecen en tu menú.")
    pagebreak(doc)

    # ─── Operations (read-only) ───
    h1(doc, "Operaciones (Vista de Lectura)")
    p(doc, "Desde acá podés consultar todas las operaciones de la agencia, pero sin poder modificarlas. Esto te permite:")
    bullet(doc, "Verificar el monto de venta y el costo de cada operación")
    bullet(doc, "Confirmar los datos para la facturación")
    bullet(doc, "Revisar los pagos asociados a cada operación")
    bullet(doc, "Chequear qué operadores están involucrados y sus costos")

    tip(doc, "Usá los filtros para buscar operaciones por estado, vendedor, agencia o rango de fechas. Esto te ayuda a encontrar rápidamente las operaciones que necesitás para conciliar.")
    pagebreak(doc)

    # ─── Operators ───
    h1(doc, "Operadores / Proveedores")
    p(doc, "Acá gestionás la información de los proveedores con los que trabaja la agencia.")

    h2(doc, "Lo que podés hacer")
    bullet(doc, "Ver la lista completa de operadores")
    bullet(doc, "Editar datos del operador (contacto, datos bancarios)")
    bullet(doc, "Consultar el balance de cada operador: cuánto se le debe, cuánto se le pagó")
    bullet(doc, "Ver las operaciones asociadas a cada operador")
    bullet(doc, "Ver los próximos vencimientos de pago")

    h2(doc, "Balance del Operador")
    p(doc, "El balance se calcula automáticamente a partir de las operaciones. Cada vez que se carga una operación con ese operador, su deuda total aumenta. Cada vez que se registra un pago, la deuda disminuye.")

    bold_p(doc, "Impacto: ", "si registrás un pago a un operador desde Caja, su balance se actualiza inmediatamente. También se genera un asiento contable automático.")
    pagebreak(doc)

    # ─── Cash ───
    h1(doc, "Caja y Bancos")
    p(doc, "Esta es una de tus secciones principales. Desde acá controlás todos los movimientos de dinero.")

    h2(doc, "Resumen")
    p(doc, "La primera pestaña te muestra:")
    bullet(doc, "Saldo actual en pesos (efectivo + banco + Mercado Pago)")
    bullet(doc, "Saldo actual en dólares")
    bullet(doc, "Totales de ingresos y egresos del período")

    h2(doc, "Movimientos de Caja")
    p(doc, "Registro de todos los movimientos de dinero. Podés filtrar por:")
    bullet(doc, "Tipo: ingreso o egreso")
    bullet(doc, "Método de pago: efectivo, banco, Mercado Pago, dólares billete")
    bullet(doc, "Fecha")
    bullet(doc, "Operación asociada")

    h2(doc, "Pagos Recibidos (Cobros de Clientes)")
    p(doc, "Desde acá ves todos los cobros pendientes y realizados. Cuando un pago se marca como \"Pagado\":")
    bullet(doc, "Se genera un asiento contable de ingreso")
    bullet(doc, "Se actualiza el saldo de caja")
    bullet(doc, "Se actualiza el estado de pago de la operación")

    h2(doc, "Egresos (Pagos a Proveedores y Otros)")
    p(doc, "Todos los pagos realizados: a operadores, comisiones, gastos generales, etc.")

    bold_p(doc, "Impacto clave: ", "marcar un pago como 'Pagado' es el disparador principal de toda la contabilidad. Ese simple clic genera el asiento contable, actualiza balances y modifica los reportes automáticamente.")

    important(doc, "Siempre verificá el método de pago correcto (efectivo, banco, MP, dólares) porque esto afecta en qué cuenta contable se registra el movimiento.")
    pagebreak(doc)

    # ─── Accounting ───
    h1(doc, "Contabilidad")
    p(doc, "La contabilidad funciona con partida doble: cada transacción genera automáticamente un débito y un crédito. La buena noticia es que casi todo se genera automáticamente desde los pagos.")

    h2(doc, "Libro Mayor")
    p(doc, "Es el registro maestro de todos los movimientos contables. Cada línea muestra:")
    bullet(doc, "Fecha del movimiento")
    bullet(doc, "Tipo (ingreso, egreso, comisión, diferencia de cambio)")
    bullet(doc, "Cuenta contable afectada")
    bullet(doc, "Monto en la moneda original")
    bullet(doc, "Monto equivalente en pesos")
    bullet(doc, "Operación asociada (si corresponde)")
    bullet(doc, "Vendedor asociado (si corresponde)")

    p(doc, "Podés filtrar por cuenta, fecha, operación o vendedor.")

    h2(doc, "Posición de IVA")
    p(doc, "Muestra mes a mes:")
    bullet(doc, "IVA débito fiscal (de las ventas)")
    bullet(doc, "IVA crédito fiscal (de las compras)")
    bullet(doc, "Saldo a pagar o a favor")

    h2(doc, "Libro de IVA")
    p(doc, "Detalle para presentar ante AFIP con todas las facturas de venta y compra del período.")

    h2(doc, "Retenciones e IIBB")
    p(doc, "Registro de retenciones de Ganancias e Ingresos Brutos, con seguimiento de pagos.")

    h2(doc, "Pagos a Operadores (Vista Contable)")
    p(doc, "Desde acá tenés una vista consolidada de todos los pagos a proveedores, organizada para la conciliación bancaria.")

    h2(doc, "Cuentas Corrientes")
    p(doc, "Muestra el saldo detallado con cada operador y cliente, con antigüedad de deuda (30, 60, 90+ días). Muy útil para saber a quién reclamar y a quién pagar.")

    h2(doc, "Manejo de Dólares y Tipo de Cambio")
    p(doc, "Cuando hay operaciones en dólares:")
    bullet(doc, "Se registra el monto original en USD")
    bullet(doc, "Se registra el equivalente en pesos al tipo de cambio del momento")
    bullet(doc, "Si más adelante el tipo de cambio varía y genera diferencia, se registra automáticamente como ganancia o pérdida por tipo de cambio")

    bold_p(doc, "Impacto: ", "toda la contabilidad se alimenta de los pagos. Si los pagos están bien cargados (monto correcto, método correcto, moneda correcta), la contabilidad va a reflejar fielmente la realidad. Si hay un error en un pago, el asiento contable también va a estar incorrecto.")
    pagebreak(doc)

    # ─── Commissions (read-only) ───
    h1(doc, "Comisiones (Vista de Lectura)")
    p(doc, "Desde acá podés consultar las comisiones de todos los vendedores, pero no podés modificarlas ni pagarlas (eso lo hace el Administrador).")

    h2(doc, "Lo que podés ver")
    bullet(doc, "Lista de comisiones por vendedor")
    bullet(doc, "Estado de cada comisión (pendiente o pagada)")
    bullet(doc, "Detalle de cálculo: operación, margen, porcentaje aplicado")
    bullet(doc, "Fecha de cálculo y fecha de pago")

    h2(doc, "Cómo se generan")
    p(doc, "Las comisiones se calculan automáticamente cuando una operación pasa a \"Confirmado\" o \"Cerrado\". El sistema toma el margen (venta menos costo) y aplica el porcentaje configurado para ese vendedor.")

    bold_p(doc, "Impacto: ", "cuando el Administrador marca una comisión como pagada, se genera un asiento contable que verás en el Libro Mayor, y se descuenta del saldo de caja.")
    pagebreak(doc)

    # ─── Reports ───
    h1(doc, "Reportes")
    p(doc, "Podés generar reportes financieros para análisis y presentación.")

    h2(doc, "Reportes que te interesan como Contable")
    bullet(doc, "Reporte de Flujo de Caja: ingresos vs. egresos por período — ideal para controlar la liquidez")
    bullet(doc, "Balance de Operadores: cuánto se le debe a cada proveedor — útil para planificar pagos")
    bullet(doc, "Antigüedad de Deuda: clientes que deben dinero, ordenados por antigüedad — para gestionar cobranzas")
    bullet(doc, "Reporte de Comisiones: resumen de comisiones por vendedor y estado")
    bullet(doc, "Reporte de Rentabilidad: margen real por operación — para verificar que los números cierran")

    p(doc, "Todos los reportes se pueden exportar a CSV o Excel.")
    pagebreak(doc)

    # ─── Alerts ───
    h1(doc, "Alertas")
    p(doc, "Las alertas que más te van a interesar son:")
    bullet(doc, "Pagos pendientes a operadores (vencimientos próximos)")
    bullet(doc, "Vencimientos impositivos (IVA, retenciones)")
    bullet(doc, "Diferencias de tipo de cambio significativas")
    bullet(doc, "Saldo bajo en caja")
    p(doc, "Podés marcar las alertas como leídas o posponerlas.")
    pagebreak(doc)

    # ─── Tasks ───
    h1(doc, "Tareas")
    p(doc, "Sistema de tareas para organizar tu trabajo diario. Podés ver tus tareas asignadas, crear nuevas y marcarlas como completadas.")
    pagebreak(doc)

    # ─── Connections ───
    h1(doc, "Cómo las Acciones se Conectan Entre Sí")
    p(doc, "Es importante entender cómo las distintas secciones se relacionan:")

    h2(doc, "Cuando se cobra un pago de cliente")
    numbered(doc, "El vendedor o administrador marca el pago como \"Pagado\"")
    numbered(doc, "Se crea automáticamente un asiento contable (ingreso)")
    numbered(doc, "Se actualiza el saldo de caja (aumenta)")
    numbered(doc, "Se actualiza el estado de la operación")
    p(doc, "Vos lo ves reflejado en: Caja (saldo), Contabilidad (asiento), y Reportes (flujo de caja)")

    h2(doc, "Cuando se paga a un operador")
    numbered(doc, "Se registra el pago en la operación o en la ficha del operador")
    numbered(doc, "Se marca como \"Pagado\"")
    numbered(doc, "Se crea el asiento contable (egreso)")
    numbered(doc, "Se actualiza el saldo de caja (disminuye)")
    numbered(doc, "Se actualiza el balance del operador (disminuye la deuda)")
    p(doc, "Vos lo ves reflejado en: Caja (saldo), Contabilidad (asiento), Operadores (balance), y Reportes")

    h2(doc, "Cuando se paga una comisión")
    numbered(doc, "El administrador marca la comisión como pagada")
    numbered(doc, "Se crea el asiento contable (egreso por comisión)")
    numbered(doc, "Se actualiza el saldo de caja (disminuye)")
    p(doc, "Vos lo ves reflejado en: Comisiones (estado cambia a pagada), Caja y Contabilidad")

    h2(doc, "Cuando hay diferencia de tipo de cambio")
    numbered(doc, "Un pago en dólares se recibe a un tipo de cambio diferente al momento de la venta")
    numbered(doc, "El sistema detecta la diferencia automáticamente")
    numbered(doc, "Se crea un asiento contable de ganancia o pérdida por tipo de cambio")
    p(doc, "Vos lo ves reflejado en: Contabilidad (asiento de diferencia de cambio), Reportes, y posiblemente en Alertas")

    doc.save(os.path.join(OUTPUT_DIR, "Manual_Contable_MaxevaGestion.docx"))
    print("✅ Manual del Contable generado")


# ─── SELLER MANUAL ────────────────────────────────────────────────────────────

def generate_seller_manual():
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc, "MAXEVA GESTIÓN", "Manual del Vendedor", "Guía para la gestión de ventas, leads y comisiones")

    h1(doc, "Contenido")
    p(doc, "1. Introducción")
    p(doc, "2. Acceso al Sistema")
    p(doc, "3. Tu Menú — Qué Podés Ver")
    p(doc, "4. Panel Principal (Tu Resumen)")
    p(doc, "5. CRM de Ventas — Tus Leads")
    p(doc, "6. Cotizaciones")
    p(doc, "7. Operaciones")
    p(doc, "8. Clientes")
    p(doc, "9. Mis Comisiones")
    p(doc, "10. Alertas")
    p(doc, "11. Documentos")
    p(doc, "12. Tareas")
    p(doc, "13. Reportes")
    p(doc, "14. Tu Flujo de Trabajo Diario")
    pagebreak(doc)

    add_intro(doc)
    p(doc, "Como Vendedor, tu foco está en las ventas: gestionar tus leads, armar cotizaciones, cerrar operaciones y seguir tus comisiones.")
    p(doc, "El sistema te muestra únicamente tu propia información: tus leads, tus operaciones, tus comisiones. No ves la información de otros vendedores ni las secciones financieras/contables.")
    pagebreak(doc)

    add_login(doc)
    pagebreak(doc)

    # ─── Menu ───
    h1(doc, "Tu Menú — Qué Podés Ver")

    h2(doc, "Secciones a las que tenés acceso")
    bullet(doc, "Resumen (Dashboard): tus propios números y métricas")
    bullet(doc, "CRM de Ventas: tus leads y oportunidades")
    bullet(doc, "Operaciones: tus operaciones confirmadas")
    bullet(doc, "Clientes: los clientes de tus operaciones")
    bullet(doc, "Alertas: avisos sobre tus viajes y pagos")
    bullet(doc, "Comisiones: lo que ganaste y lo que te deben")
    bullet(doc, "Documentos: documentos de tus operaciones")
    bullet(doc, "Reportes: tus reportes de ventas y rendimiento")
    bullet(doc, "Tareas: tus tareas pendientes")

    h2(doc, "Secciones que NO ves")
    bullet(doc, "Operadores / Proveedores")
    bullet(doc, "Caja y Bancos")
    bullet(doc, "Contabilidad")
    bullet(doc, "Configuración")
    p(doc, "Estas secciones son de uso administrativo y contable.")

    important(doc, "Solo ves TU información. No podés ver los leads, operaciones o comisiones de otros vendedores.")
    pagebreak(doc)

    # ─── Dashboard ───
    h1(doc, "Panel Principal (Tu Resumen)")
    p(doc, "Al entrar al sistema, lo primero que ves es tu resumen personal con tus números del período:")
    bullet(doc, "Cuánto vendiste (en pesos y dólares)")
    bullet(doc, "Cuántas operaciones tenés activas")
    bullet(doc, "Comisiones pendientes de cobro")
    bullet(doc, "Gráficos de tus ventas por destino")

    tip(doc, "Revisá tu Dashboard todos los días para tener una foto clara de cómo vas en el mes.")
    pagebreak(doc)

    # ─── CRM ───
    h1(doc, "CRM de Ventas — Tus Leads")
    p(doc, "El CRM es tu herramienta principal de trabajo. Acá gestionás todas las consultas y oportunidades de venta.")

    h2(doc, "Cómo llegan tus leads")
    numbered(doc, "Los creás vos manualmente con el botón \"Nuevo Lead\"")
    numbered(doc, "Llegan automáticamente desde Trello (si tu agencia usa esta integración)")
    numbered(doc, "Llegan desde Manychat/WhatsApp cuando un contacto nuevo escribe")
    numbered(doc, "Podés reclamar leads sin asignar que aparecen en la columna \"Sin Asignar\"")

    h2(doc, "Vista Kanban")
    p(doc, "El tablero Kanban te muestra tus leads organizados en columnas:")
    bullet(doc, "Nuevo: consultas que acaban de llegar")
    bullet(doc, "En Progreso: estás buscando opciones para el cliente")
    bullet(doc, "Cotizado: ya le mandaste una propuesta")
    bullet(doc, "Ganado: el cliente confirmó (se convierte en operación)")
    bullet(doc, "Perdido: el cliente no avanzó")

    p(doc, "Para cambiar el estado de un lead, simplemente arrastrá la tarjeta a otra columna.")

    h2(doc, "Crear un Lead Nuevo")
    numbered(doc, "Hacé clic en \"Nuevo Lead\"")
    numbered(doc, "Completá el nombre del contacto")
    numbered(doc, "Poné su teléfono (se va a crear un botón directo a WhatsApp)")
    numbered(doc, "Indicá el destino de interés")
    numbered(doc, "Guardá")

    h2(doc, "Detalle del Lead")
    p(doc, "Hacé clic en cualquier lead para abrir su ficha. Desde ahí podés:")
    bullet(doc, "Ver y editar los datos del contacto")
    bullet(doc, "Crear una cotización para enviarle al cliente")
    bullet(doc, "Ver cotizaciones anteriores")
    bullet(doc, "Convertir el lead en operación cuando el cliente confirma")
    bullet(doc, "Contactar al cliente por WhatsApp con un clic")

    bold_p(doc, "Impacto: ", "cuando movés un lead a \"Ganado\" y lo convertís en operación, se crean automáticamente la operación y el registro del cliente. El lead desaparece del CRM activo y pasa al historial.")

    tip(doc, "Mantené tus leads actualizados. Mové a \"Perdido\" los que no van a avanzar para que tu tablero refleje solo oportunidades reales.")
    pagebreak(doc)

    # ─── Quotations ───
    h1(doc, "Cotizaciones")
    p(doc, "Las cotizaciones son propuestas de viaje que le armás al cliente para que vea opciones y precios.")

    h2(doc, "Crear una Cotización")
    numbered(doc, "Abrí el detalle de un lead")
    numbered(doc, "Hacé clic en \"Cotizar\"")
    numbered(doc, "Se abre el armador de cotización")

    h2(doc, "Agregar Servicios")
    p(doc, "Podés agregar diferentes tipos de servicio:")

    h3(doc, "Hotel")
    p(doc, "Escribí el nombre del hotel y el sistema te va a sugerir opciones de una base de más de 1.600 hoteles. Al seleccionar uno:")
    bullet(doc, "Se completan automáticamente las estrellas")
    bullet(doc, "Se agrega la dirección (desde Google)")
    bullet(doc, "Se muestra una foto del hotel")
    p(doc, "Después completá: tipo de habitación, régimen de comidas, fechas de check-in/check-out y cantidad de noches.")

    h3(doc, "Vuelo")
    p(doc, "Indicá la aerolínea, ruta (ej: Buenos Aires - Cancún), fechas de ida y vuelta, clase (turista, business) y si tiene escalas.")

    h3(doc, "Asistencia al Viajero")
    p(doc, "Agregá el seguro de viaje indicando el proveedor, tipo de cobertura y precio.")

    h3(doc, "Traslado")
    p(doc, "Agregá transfers (ej: \"Aeropuerto - Hotel - Aeropuerto\") con su descripción y precio.")

    h3(doc, "Otro / Paquete")
    p(doc, "Para cualquier servicio adicional o para armar un paquete que combine todo.")

    h2(doc, "Opciones Múltiples")
    p(doc, "Podés armar hasta 3 opciones dentro de la misma cotización. Esto es ideal para ofrecer alternativas al cliente (ej: hotel 4 estrellas vs 5 estrellas).")

    h2(doc, "Compartir con el Cliente")
    p(doc, "Una vez guardada la cotización, se genera un link que podés copiar y enviar por WhatsApp o email. El cliente abre ese link y ve una página profesional con:")
    bullet(doc, "Todos los servicios incluidos")
    bullet(doc, "Fotos de los hoteles")
    bullet(doc, "Precios detallados")
    bullet(doc, "El total del viaje")

    bold_p(doc, "Impacto: ", "cuando creás una cotización, el lead pasa automáticamente a estado \"Cotizado\". La cotización queda guardada en el historial del lead.")
    pagebreak(doc)

    # ─── Operations ───
    h1(doc, "Operaciones")
    p(doc, "Una operación es un viaje confirmado. Acá ves todas las operaciones que gestionás.")

    h2(doc, "Cómo se Crea una Operación")
    p(doc, "La forma más común es convertir un lead ganado en operación. El sistema crea todo automáticamente con los datos del lead.")

    h2(doc, "Lo que ves en cada Operación")
    bullet(doc, "Datos del viaje: destino, fechas, tipo de operación")
    bullet(doc, "Pasajeros: quiénes viajan")
    bullet(doc, "Precio de venta")
    bullet(doc, "Estado de la operación (Pre-reserva, Reservado, Confirmado, etc.)")
    bullet(doc, "Pagos del cliente: cuánto pagó y cuánto falta")

    h2(doc, "Estados de la Operación")
    bullet(doc, "Pre-reserva: estás reservando con el proveedor")
    bullet(doc, "Reservado: reserva confirmada con el proveedor")
    bullet(doc, "Confirmado: todo pago, viaje listo")
    bullet(doc, "Viajado: el cliente ya viajó")
    bullet(doc, "Cerrado: todo liquidado")

    bold_p(doc, "Impacto: ", "cuando la operación pasa a \"Confirmado\", se calcula automáticamente tu comisión. Cuando pasa a \"Cerrado\", se liquida todo.")

    tip(doc, "Mantené actualizados los estados de tus operaciones. Esto no solo ordena tu trabajo, sino que dispara las comisiones y las alertas automáticas.")
    pagebreak(doc)

    # ─── Customers ───
    h1(doc, "Clientes")
    p(doc, "Acá ves los clientes asociados a tus operaciones.")

    h2(doc, "Datos del Cliente")
    bullet(doc, "Nombre y apellido")
    bullet(doc, "Email y teléfono")
    bullet(doc, "Fecha de nacimiento")
    bullet(doc, "Tipo y número de documento")

    h2(doc, "Subir Documentos de Pasajeros")
    p(doc, "Podés subir una foto del pasaporte o DNI del cliente. El sistema lee automáticamente los datos (nombre, número de documento, fecha de nacimiento) y te los muestra para confirmar.")

    bold_p(doc, "Impacto: ", "los clientes se crean automáticamente cuando convertís un lead en operación. Si subís un pasaporte, los datos se actualizan sin que tengas que tipearlos.")
    pagebreak(doc)

    # ─── Commissions ───
    h1(doc, "Mis Comisiones")
    p(doc, "Esta es probablemente tu sección favorita. Acá ves todo lo que ganaste y lo que te deben.")

    h2(doc, "Cómo se Calcula tu Comisión")
    p(doc, "Por cada operación que cerrás:")
    numbered(doc, "El sistema toma el precio de venta del viaje")
    numbered(doc, "Le resta el costo del proveedor (lo que le paga la agencia al operador)")
    numbered(doc, "La diferencia es el margen de ganancia")
    numbered(doc, "Tu comisión es un porcentaje de ese margen")

    p(doc, "Por ejemplo: vendiste un viaje a $1.000 y el operador cobra $700. El margen es $300. Si tu porcentaje es 30%, tu comisión es $90.")

    h2(doc, "Vendedor Secundario")
    p(doc, "Si en una operación hay un vendedor principal y uno secundario, la comisión se divide entre los dos (generalmente 50/50).")

    h2(doc, "Lo que ves")
    bullet(doc, "Lista de todas tus comisiones")
    bullet(doc, "Estado de cada una: Pendiente (te la deben) o Pagada")
    bullet(doc, "El detalle de cálculo: operación, margen, porcentaje, monto")
    bullet(doc, "Total acumulado pendiente de cobro")

    h2(doc, "Mi Balance")
    p(doc, "Desde la sección \"Mi Balance\" ves un resumen de:")
    bullet(doc, "Total de comisiones generadas (en pesos y dólares)")
    bullet(doc, "Total ya cobrado")
    bullet(doc, "Total pendiente de cobro")

    bold_p(doc, "Impacto: ", "las comisiones se calculan automáticamente cuando tu operación llega a \"Confirmado\". No necesitás hacer nada manual. El Administrador es quien marca las comisiones como pagadas.")

    important(doc, "Si querés que tus comisiones se calculen, asegurate de que la operación tenga cargado correctamente: el precio de venta, el operador con su costo, y que el estado esté en \"Confirmado\" o \"Cerrado\".")
    pagebreak(doc)

    # ─── Alerts ───
    h1(doc, "Alertas")
    p(doc, "El sistema te avisa automáticamente de cosas importantes:")
    bullet(doc, "Viaje próximo: 48 a 72 horas antes de la salida de un viaje tuyo")
    bullet(doc, "Pago pendiente: cuando un cliente tiene un pago vencido")
    bullet(doc, "Documentación faltante: cuando un pasajero no tiene documentos cargados")

    p(doc, "Las alertas aparecen como notificaciones. Podés marcarlas como leídas o posponerlas.")

    tip(doc, "Revisá tus alertas al inicio de cada día de trabajo. Te ayudan a no olvidar seguimientos importantes.")
    pagebreak(doc)

    # ─── Documents ───
    h1(doc, "Documentos")
    p(doc, "Acá podés ver y subir documentos relacionados con tus operaciones: pasaportes, vouchers, comprobantes de pago, etc.")

    h2(doc, "Subir un Documento")
    numbered(doc, "Abrí la operación correspondiente")
    numbered(doc, "Andá a la pestaña de Documentos")
    numbered(doc, "Hacé clic en \"Subir Documento\"")
    numbered(doc, "Seleccioná el archivo (foto, PDF)")
    numbered(doc, "Si es un pasaporte o DNI, el sistema va a leer los datos automáticamente")
    pagebreak(doc)

    # ─── Tasks ───
    h1(doc, "Tareas")
    p(doc, "Podés crear tareas personales para organizar tu trabajo:")
    bullet(doc, "\"Llamar a Juan por el viaje a Cancún\"")
    bullet(doc, "\"Enviar cotización a María\"")
    bullet(doc, "\"Confirmar reserva con operador\"")
    p(doc, "Cada tarea tiene fecha de vencimiento y podés marcarla como completada.")
    pagebreak(doc)

    # ─── Reports ───
    h1(doc, "Reportes")
    p(doc, "Podés ver tus propios reportes:")
    bullet(doc, "Reporte de Ventas: cuánto vendiste por período y por destino")
    bullet(doc, "Reporte de Comisiones: detalle de tus comisiones por operación")
    p(doc, "Podés exportarlos a Excel para llevar tu propio seguimiento.")
    pagebreak(doc)

    # ─── Daily flow ───
    h1(doc, "Tu Flujo de Trabajo Diario")
    p(doc, "Una guía rápida de cómo usar el sistema en tu día a día:")

    h2(doc, "Al empezar el día")
    numbered(doc, "Revisá tu Dashboard para ver cómo van tus números")
    numbered(doc, "Mirá las Alertas por si hay algo urgente (viaje mañana, pago vencido)")
    numbered(doc, "Chequeá tus Tareas pendientes")

    h2(doc, "Cuando llega una consulta")
    numbered(doc, "Creá un nuevo Lead en el CRM (o reclamá uno sin asignar)")
    numbered(doc, "Completá los datos del contacto y el destino de interés")
    numbered(doc, "Empezá a buscar opciones")

    h2(doc, "Cuando tenés opciones listas")
    numbered(doc, "Abrí el lead y creá una Cotización")
    numbered(doc, "Agregá todos los servicios: hotel, vuelo, asistencia, traslados")
    numbered(doc, "Guardá y compartí el link con el cliente por WhatsApp")
    numbered(doc, "Mové el lead a \"Cotizado\"")

    h2(doc, "Cuando el cliente confirma")
    numbered(doc, "Abrí el lead y hacé clic en \"Convertir en Operación\"")
    numbered(doc, "El sistema crea la operación y el cliente automáticamente")
    numbered(doc, "Completá los datos del operador y el costo")
    numbered(doc, "Registrá los pagos del cliente")

    h2(doc, "Seguimiento de la operación")
    numbered(doc, "Actualizá el estado de la operación a medida que avanza")
    numbered(doc, "Cuando llegue a \"Confirmado\", verificá que tu comisión se haya calculado en \"Mis Comisiones\"")
    numbered(doc, "Subí los documentos de los pasajeros cuando los tengas")

    h2(doc, "Al final del mes")
    numbered(doc, "Revisá tu balance de comisiones")
    numbered(doc, "Generá tu reporte de ventas")
    numbered(doc, "Revisá qué leads quedaron sin cerrar y hacé seguimiento")

    doc.save(os.path.join(OUTPUT_DIR, "Manual_Vendedor_MaxevaGestion.docx"))
    print("✅ Manual del Vendedor generado")


# ─── Main ────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    generate_admin_manual()
    generate_contable_manual()
    generate_seller_manual()
    print(f"\n📁 Manuales guardados en: {OUTPUT_DIR}")
